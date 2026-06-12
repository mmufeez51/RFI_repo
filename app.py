import streamlit as st
import google.generativeai as genai
from docx import Document
import io
import os
from PIL import Image

# Configure page
st.set_page_config(page_title="AI RFI Assistant", page_icon="🏗️", layout="wide")

# Inject Custom CSS for an elegant background
elegant_css = """
<style>
    /* Main app background gradient */
    .stApp {
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
    }
    
    /* Elegant translucent sidebar */
    [data-testid="stSidebar"] {
        background-color: rgba(255, 255, 255, 0.4) !important;
        backdrop-filter: blur(10px) !important;
    }
    
    /* Slightly transparent inputs for a glassmorphism feel */
    div.stTextInput > div > div > input,
    div.stTextArea > div > div > textarea {
        background-color: rgba(255, 255, 255, 0.7) !important;
        border: 1px solid rgba(255, 255, 255, 0.5) !important;
        border-radius: 8px !important;
    }
    
    /* Primary button premium styling */
    button[kind="primary"] {
        background: linear-gradient(to right, #4facfe 0%, #00f2fe 100%) !important;
        border: none !important;
        box-shadow: 0 4px 15px rgba(0, 242, 254, 0.4) !important;
        color: white !important;
        border-radius: 8px !important;
    }
</style>
"""
st.markdown(elegant_css, unsafe_allow_html=True)

# Set up the API key for Google Gemini
# Prioritize Streamlit secrets, fallback to environment variable
API_KEY = None
try:
    if "GEMINI_API_KEY" in st.secrets:
        API_KEY = st.secrets["GEMINI_API_KEY"]
except Exception:
    pass

if not API_KEY:
    API_KEY = os.environ.get("GEMINI_API_KEY")

if API_KEY:
    genai.configure(api_key=API_KEY)
else:
    st.warning("⚠️ Please set your GEMINI_API_KEY as an environment variable or in Streamlit secrets (.streamlit/secrets.toml) to use the AI features.")

# Sidebar Configuration
st.sidebar.header("⚙️ Configuration")
rfi_tone = st.sidebar.selectbox("Tone", ["Professional", "Urgent", "Friendly", "Highly Formal"])
target_audience = st.sidebar.selectbox("Target Audience", ["Architect", "Structural Engineer", "Client", "Subcontractor", "General"])

st.title("🏗️ AI-Powered RFI Assistant")
st.write("Upload a construction photo, provide some project details, and the AI will generate a formal Request for Information (RFI) document for you to edit and download.")

# Initialize session state for the draft
if 'rfi_draft' not in st.session_state:
    st.session_state.rfi_draft = ""

# Setup Tabs
tab1, tab2 = st.tabs(["1. Setup & Generate 🏗️", "2. Review & Edit 📝"])

with tab1:
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.subheader("Project Details")
        project_name = st.text_input("Project Name", placeholder="e.g., Downtown Plaza Revitalization")
        rfi_number = st.text_input("RFI Number", placeholder="e.g., RFI-042")
        requested_by = st.text_input("Requested By", placeholder="e.g., John Doe - Project Manager")
        
        site_notes = st.text_area("Site Notes", placeholder="Enter your observations, measurements, or specific queries here...", height=150)
        
    with col2:
        st.subheader("Context Image")
        uploaded_file = st.file_uploader("Upload Construction Photo", type=["jpg", "jpeg", "png"])
        if uploaded_file is not None:
            st.image(uploaded_file, caption="Uploaded Construction Photo", use_column_width=True)

    if st.button("✨ Generate RFI Draft", type="primary", use_container_width=True):
        if not API_KEY:
            st.error("Gemini API Key is missing. Cannot generate RFI. Please set the GEMINI_API_KEY environment variable.")
        elif not uploaded_file and not site_notes:
            st.error("Please upload an image and/or provide site notes to generate an RFI.")
        else:
            with st.spinner("Analyzing image and notes to generate formal RFI..."):
                try:
                    # We use the requested 'gemini-3.5-flash' model
                    model = genai.GenerativeModel('gemini-3.5-flash')
                    
                    contents = []
                    
                    # Instruction prompt
                    prompt = f"""
                    You are an expert construction project manager. Based on the provided construction site photo and the site notes, please generate a structured, formal Request for Information (RFI). 
                    
                    Configuration:
                    - Tone: {rfi_tone}
                    - Target Audience: {target_audience}
                    - Project Name: {project_name if project_name else 'Not provided'}
                    - RFI Number: {rfi_number if rfi_number else 'Not provided'}
                    - Requested By: {requested_by if requested_by else 'Not provided'}
                    
                    The RFI should be formatted clearly and include the following sections:
                    - Subject Line (clear and concise)
                    - Background / Context (based on the image and notes provided)
                    - The Specific Question or Clarification Needed
                    - Potential Impact if not resolved (e.g., schedule delay, cost impact)
                    - Suggested Solution (if any can be inferred)
                    
                    Make the tone appropriate based on the configuration above, and ready to be sent to the target audience. Do not include placeholder brackets like [Insert Date], just provide the core content.
                    """
                    contents.append(prompt)
                    
                    if uploaded_file is not None:
                        image = Image.open(uploaded_file)
                        contents.append(image)
                        
                    if site_notes:
                        contents.append(f"Site Notes from field personnel:\n{site_notes}")

                    # Generate content
                    response = model.generate_content(contents)
                    
                    # Store in session state
                    st.session_state.rfi_draft = response.text
                    st.success("✅ RFI Generated Successfully! Go to the '2. Review & Edit 📝' tab to finalize it.")
                    
                except Exception as e:
                    st.error(f"An error occurred during generation: {e}")

def create_docx(text):
    doc = Document()
    doc.add_heading('Request for Information (RFI)', 0)
    
    # Simple parsing to preserve some formatting
    paragraphs = text.split('\n')
    for para in paragraphs:
        if para.strip() != "":
            # Handle basic bolding and headers for cleaner word doc output
            if para.startswith('**') and para.endswith('**'):
                p = doc.add_paragraph()
                p.add_run(para.replace('**', '')).bold = True
            elif para.startswith('#'):
                doc.add_heading(para.replace('#', '').strip(), level=2)
            else:
                doc.add_paragraph(para)
                
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

with tab2:
    st.subheader("Edit Your Draft")
    if not st.session_state.rfi_draft:
        st.info("Your generated RFI draft will appear here. Please go to the first tab to generate one.")
    else:
        st.write("You can freely edit the text below. Once you are happy with the wording, download the final document.")
        edited_rfi = st.text_area("RFI Content", value=st.session_state.rfi_draft, height=400, label_visibility="collapsed")
        
        st.markdown("---")
        
        # Create docx for download from the EDITED text
        docx_buffer = create_docx(edited_rfi)
        
        st.download_button(
            label="📄 Download Final RFI as Word Document (.docx)",
            data=docx_buffer,
            file_name=f"RFI_{rfi_number if rfi_number else 'Draft'}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary"
        )
